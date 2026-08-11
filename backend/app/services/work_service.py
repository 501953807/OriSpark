"""作品管理业务服务 — 对应: docs/modules-v5/01-creative-assets.md
Phase 1.3: 视频缩略图30%位置关键帧 (非首帧)
增强: 失败时记录日志，不再静默吞错误"""

import logging
import os
import re
import subprocess
from pathlib import Path
from typing import Optional, Tuple, Dict, Any

logger = logging.getLogger("oristudio.works")

# 文件存储根目录
UPLOAD_DIR = Path("data/workspace")
THUMBNAIL_DIR = Path("data/thumbnails")

# 支持的文件类型 (P1.7.13: dictStore-backed with hardcoded fallback)
ALLOWED_EXTENSIONS = {
    "jpg", "jpeg", "png", "webp", "gif", "svg", "bmp", "tiff",
    "mp3", "wav", "flac", "ogg", "aac", "m4a",
    "mp4", "mov", "webm", "avi", "mkv",
    "pdf", "docx", "doc", "txt", "md", "rtf",
    "psd", "ai", "fig", "sketch",
    "py", "js", "ts", "html", "css", "json", "zip",
}

# P2-1: RAW camera image extensions (case-insensitive matching)
RAW_EXTENSIONS = {
    "cr2", "cr3", "nef", "arw", "raf", "orf", "pef", "dng", "heic", "heif",
    "rw2", "x3f", "iiq", "sr2", "mos", "mef", "k25", "kdc", "srf", "bay", "ptx", "dcraw", "raw",
}

MAX_FILE_SIZE = 500 * 1024 * 1024  # 500MB


def _get_allowed_extensions(db) -> set:
    """Get allowed file extensions (dictStore-backed, P1.7.13)."""
    try:
        from app.utils.system_helpers import get_dict_values
        dict_exts = get_dict_values("file_extensions", db)
        if dict_exts:
            return set(dict_exts)
    except Exception as e:
        logger.exception("Error in _get_allowed_extensions: %s", str(e))
    return ALLOWED_EXTENSIONS


def get_image_dimensions(file_path: str) -> Tuple[Optional[int], Optional[int]]:
    try:
        from PIL import Image
        img = Image.open(file_path)
        return img.size
    except Exception:
        return None, None


def detect_file_type(extension: str, content: bytes = b"") -> str:
    ext = extension.lower()
    # P2-1: RAW format support
    raw_exts = {"cr2","nef","arw","dng","rw2","orf","pef","raf","x3f","iiq","sr2","mos","mef","k25","kdc","srf","bay","ptx","dcraw","raw"}
    if ext in raw_exts: return "image"  # RAW images are still images
    if ext in {"jpg","jpeg","png","webp","gif","svg","bmp","tiff"}: return "image"
    if ext in {"mp3","wav","flac","ogg","aac","m4a"}: return "audio"
    if ext in {"mp4","mov","webm","avi","mkv"}: return "video"
    if ext in {"pdf","docx","doc","txt","md","rtf"}: return "document"
    if ext in {"psd","ai","fig","sketch"}: return "design"
    if ext in {"py","js","ts","html","css","json","xml","yaml","zip","tar","gz"}: return "code"
    return "other"


def generate_thumbnail(file_path: str, file_type: str, work_id: str) -> Optional[str]:
    thumb_dir = Path("data/thumbnails") / work_id[:2]
    thumb_dir.mkdir(parents=True, exist_ok=True)
    thumb_path = thumb_dir / f"{work_id}_thumb.jpg"

    try:
        if file_type == "image":
            from PIL import Image
            img = Image.open(file_path)
            img.thumbnail((400, 400), Image.LANCZOS)
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            img.save(thumb_path, "JPEG", quality=85)
            logger.debug(f"Thumbnail generated: {file_path} -> {thumb_path}")
            return str(thumb_path.resolve())

        elif file_type == "video":
            # Phase 1.3: 提取30%位置关键帧 (非首帧)
            try:
                probe = subprocess.run([
                    "ffprobe", "-v", "error", "-show_entries", "format=duration",
                    "-of", "default=noprint_wrappers=1:nokey=1", file_path
                ], capture_output=True, text=True, timeout=15)
                duration = float(probe.stdout.strip()) if probe.stdout.strip() else 0
            except Exception:
                logger.warning(f"ffprobe failed for {file_path}, using duration=0")
                duration = 0
            seek_time = duration * 0.3 if duration > 3 else (duration * 0.5 if duration > 0 else 0)
            r = subprocess.run([
                "ffmpeg", "-ss", str(seek_time), "-i", file_path,
                "-vframes", "1", "-s", "400x300", "-f", "image2",
                str(thumb_path), "-y"
            ], capture_output=True, timeout=30)
            if r.returncode == 0 and thumb_path.exists():
                logger.debug(f"Video thumbnail generated: {file_path} at {seek_time}s")
                return str(thumb_path.resolve())
            else:
                logger.warning(f"ffmpeg thumbnail failed for {file_path}: {r.stderr.decode()[:200] if r.stderr else 'unknown error'}")

        elif file_type == "audio":
            from PIL import Image, ImageDraw
            img = Image.new("RGB", (400, 300), (240, 240, 245))
            draw = ImageDraw.Draw(img)
            import random
            random.seed(hash(file_path) % 10000)
            for i in range(0, 400, 4):
                h = random.randint(20, 200)
                draw.rectangle([i, 150 - h // 2, i + 2, 150 + h // 2], fill=(100, 180, 160))
            img.save(thumb_path, "JPEG", quality=85)
            return str(thumb_path.resolve())

        elif file_type == "document":
            from PIL import Image, ImageDraw
            img = Image.new("RGB", (400, 300), (245, 245, 250))
            draw = ImageDraw.Draw(img)
            draw.text((200, 150), "DOC", fill=(100, 100, 100), anchor="mm")
            img.save(thumb_path, "JPEG", quality=85)
            return str(thumb_path.resolve())

        elif file_type == "design":
            from PIL import Image, ImageDraw
            img = Image.new("RGB", (400, 300), (250, 245, 240))
            draw = ImageDraw.Draw(img)
            draw.text((200, 150), "DESIGN", fill=(100, 100, 100), anchor="mm")
            img.save(thumb_path, "JPEG", quality=85)
            return str(thumb_path.resolve())

        else:
            # Generic fallback for any file type
            from PIL import Image, ImageDraw
            img = Image.new("RGB", (400, 300), (245, 245, 250))
            draw = ImageDraw.Draw(img)
            draw.text((200, 150), file_type.upper()[:8], fill=(100, 100, 100), anchor="mm")
            img.save(thumb_path, "JPEG", quality=85)
            return str(thumb_path.resolve())

    except FileNotFoundError:
        logger.warning(f"File not found for thumbnail: {file_path}")
        return None
    except Exception as exc:
        logger.error(f"Thumbnail generation failed for {file_path}: {exc}")
        return None


def extract_exif(file_path: str) -> Optional[dict]:
    """提取图片EXIF元数据，返回键名与works_router查询条件兼容的字典."""
    try:
        from PIL import Image
        from PIL.ExifTags import TAGS
        img = Image.open(file_path)
        data = img.getexif()
        if not data:
            return None
        result = {}
        # Map standard EXIF keys to query-expectation names matching works_router
        key_mapping = {
            "Make": "CameraMake",
            "Model": "CameraModel",
            "ISOSpeedRatings": "ISOSpeed",
            "DateTimeOriginal": "DateTimeOriginal",
            "GPSLatitude": "GPSLatitude",
            "GPSLongitude": "GPSLongitude",
        }
        for tid, val in data.items():
            name = TAGS.get(tid, tid)
            final_name = key_mapping.get(name, name)
            if isinstance(val, bytes):
                try:
                    val = val.decode("utf-8", errors="replace")
                except UnicodeDecodeError:
                    val = val.hex()
            result[final_name] = str(val)
        return result
    except Exception:
        return None


# ====== 新增元数据提取 ======

def extract_audio_metadata(file_path: str) -> Dict[str, Any]:
    """使用 mutagen 提取音频元数据: 时长/采样率/比特率/艺术家/专辑."""
    meta: Dict[str, Any] = {}
    try:
        from mutagen import File as MutagenFile
        audio = MutagenFile(file_path)
        if audio is None: return meta
        if hasattr(audio.info, 'length'):
            meta["duration"] = round(audio.info.length, 1)
        if hasattr(audio.info, 'sample_rate'):
            meta["sample_rate"] = audio.info.sample_rate
        if hasattr(audio.info, 'bitrate'):
            meta["bitrate"] = audio.info.bitrate
        # tags
        for key in ['artist','album','title','genre','date']:
            if key in audio:
                meta[key] = str(audio[key][0]) if audio[key] else None
    except ImportError:
        pass  # mutagen 未安装
    except Exception:
        pass
    return meta


def extract_video_metadata(file_path: str) -> Dict[str, Any]:
    """使用 ffprobe 提取视频元数据: 分辨率/帧率/编码/时长."""
    meta: Dict[str, Any] = {}
    try:
        r = subprocess.run([
            "ffprobe","-v","quiet","-print_format","json","-show_format","-show_streams",file_path
        ], capture_output=True, text=True, timeout=30)
        if r.returncode != 0: return meta
        import json
        data = json.loads(r.stdout)
        for stream in data.get("streams",[]):
            if stream.get("codec_type")=="video":
                meta["width"] = stream.get("width")
                meta["height"] = stream.get("height")
                meta["codec"] = stream.get("codec_name")
                meta["fps"] = eval(stream.get("r_frame_rate","0/1"))
            elif stream.get("codec_type")=="audio":
                meta["audio_codec"] = stream.get("codec_name")
        fmt = data.get("format",{})
        meta["duration"] = round(float(fmt.get("duration",0)),1) if fmt.get("duration") else None
        meta["bitrate"] = int(fmt["bitrate"])//1000 if fmt.get("bitrate") else None
    except Exception:
        pass
    return meta


def extract_document_metadata(file_path: str) -> Dict[str, Any]:
    """提取文档元数据: 字数/页数/语言."""
    meta: Dict[str, Any] = {}
    ext = Path(file_path).suffix.lower()
    try:
        if ext == ".txt" or ext == ".md":
            with open(file_path,"r",encoding="utf-8",errors="replace") as f:
                text = f.read()
            meta["char_count"] = len(text)
            meta["word_count"] = len(text.split())
            meta["line_count"] = text.count("\n")+1
        elif ext == ".pdf":
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                meta["pages"] = len(reader.pages)
            except ImportError:
                pass
        elif ext == ".docx":
            try:
                from docx import Document
                doc = Document(file_path)
                meta["paragraphs"] = len(doc.paragraphs)
                text = " ".join(p.text for p in doc.paragraphs)
                meta["char_count"] = len(text)
            except ImportError:
                pass
    except Exception:
        pass
    return meta


def get_all_metadata(file_path: str, file_type: str) -> Dict[str, Any]:
    """根据文件类型提取所有元数据."""
    meta: Dict[str, Any] = {}
    if file_type == "image":
        dims = get_image_dimensions(file_path)
        if dims[0]: meta["width"], meta["height"] = dims
        exif = extract_exif(file_path)
        if exif: meta["exif_data"] = exif
    elif file_type == "audio":
        meta.update(extract_audio_metadata(file_path))
    elif file_type == "video":
        meta.update(extract_video_metadata(file_path))
    elif file_type == "document":
        meta.update(extract_document_metadata(file_path))
    return meta


# ==============================
# Security sanitization helpers
# ==============================

def sanitize_tag(tag: str) -> str:
    """Sanitize a tag string by removing HTML tags, control characters, and limiting length."""
    if not tag or not isinstance(tag, str):
        return ""
    clean_tag = re.sub(r'<[^>]*>', '', tag)
    clean_tag = re.sub(r'[\x00-\x1f\x7f]', '', clean_tag)
    clean_tag = clean_tag.strip()
    clean_tag = clean_tag[:100]
    return clean_tag


def sanitize_filename(filename: str) -> str:
    """Sanitize a filename by removing dangerous characters and path components."""
    if not filename or not isinstance(filename, str):
        return "uploaded_file"
    filename = os.path.basename(filename)
    filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
    filename = re.sub(r'[\x00-\x1f\x7f]', '', filename)
    filename = filename.strip()
    if not filename:
        filename = "uploaded_file"
    return filename


# ═══════════════════════════════════════════
# Phase 1.1: 自动元数据提取辅助函数
# ═══════════════════════════════════════════

def _extract_title_from_filename(filename: str) -> str:
    """Extract title from filename with sanitization."""
    cleaned_name = sanitize_filename(filename)
    name = os.path.splitext(cleaned_name)[0]
    name = re.sub(r'^[\d\-_]{6,}', '', name).strip('_ ')
    name = re.sub(r'^(IMG|DSC|PXL|DSCF|MVI|VID)_?\d*[_\-]?', '', name, flags=re.IGNORECASE).strip('_ ')
    return name or "未命名作品"


def _extract_completion_date(exif_data, file_path: str):
    if exif_data:
        for key in ("DateTimeOriginal", "DateTimeDigitized", "DateTime"):
            if exif_data.get(key):
                dt = str(exif_data[key])
                parts = dt.replace(" ", ":").split(":")
                if len(parts) >= 3:
                    return f"{parts[0]}-{parts[1]}-{parts[2]}"
                return dt
    try:
        from datetime import datetime
        return datetime.fromtimestamp(os.path.getmtime(file_path)).strftime("%Y-%m-%d")
    except OSError:
        return None


def _extract_creation_tool(exif_data, full_meta: dict):
    if exif_data:
        for key in ("Software", "HostComputer", "ProcessingSoftware"):
            if exif_data.get(key):
                return str(exif_data[key])
    if full_meta:
        for key in ("encoder", "producer", "software"):
            if full_meta.get(key):
                return str(full_meta[key])
    return None


def _extract_creation_location(exif_data):
    if not exif_data:
        return None
    gps_lat = exif_data.get("GPSLatitude") or exif_data.get("GPSInfo")
    gps_lon = exif_data.get("GPSLongitude")
    if gps_lat and gps_lon:
        return f"{gps_lat}, {gps_lon}"
    return None


def _build_auto_rights(exif_data):
    rights: dict = {}
    if exif_data:
        if exif_data.get("Artist"):
            rights["author_name"] = str(exif_data["Artist"])
        if exif_data.get("Copyright"):
            rights["copyright_year"] = str(exif_data["Copyright"])[:30]
    return rights


def _detect_creator_type(file_type: str, exif_data: Optional[dict], full_meta: Optional[dict]) -> str:
    """Detect creator type from file characteristics."""
    if file_type == "image" and exif_data:
        camera_signals = ("CameraMake", "CameraModel", "LensModel", "Model")
        if any(exif_data.get(k) for k in camera_signals):
            return "photographer"
    if file_type == "audio":
        return "musician"
    if file_type == "video":
        return "video"
    if file_type == "document":
        if full_meta and any(k in full_meta for k in ("software", "Application", "Producer")):
            app = str(full_meta.get("software", "") + full_meta.get("Application", "") + full_meta.get("Producer", ""))
            if any(kw in app.lower() for kw in ("autocad", "solidworks", "sketchup")):
                return "craftsman"
        return "writer"
    if file_type == "design":
        return "illustrator"
    return "illustrator"


def _thumb_to_api_path(thumb_path: Optional[str]) -> Optional[str]:
    """将本地缩略图绝对路径转为 API 可访问的相对路径."""
    if not thumb_path:
        return None
    try:
        p = Path(thumb_path)
        rel = p.relative_to(Path("data").resolve())
        return f"/api/files/{rel.as_posix()}"
    except (ValueError, OSError):
        return None


def _work_to_response(work) -> dict:
    """将 Work ORM 对象转为前端友好的响应格式 (Phase 1.1: 含自动元数据)."""
    from app.schemas.work import WorkResponse
    data = WorkResponse.model_validate(work).model_dump()
    if work.file_path:
        try:
            rel = Path(work.file_path).relative_to(Path("data").resolve())
            data["file_url"] = f"/api/files/{rel.as_posix()}"
        except (ValueError, OSError):
            data["file_url"] = None
    else:
        data["file_url"] = None

    data["thumbnail_url"] = _thumb_to_api_path(work.thumbnail_path)

    cm = work.custom_metadata or {}
    for key in ("completion_date", "creation_tool", "creation_location", "auto_tags"):
        if cm.get(key) and key not in data:
            data[key] = cm[key]

    data["verified_status"] = "已存证 ✅" if work.is_verified else None

    data["is_raw_original"] = work.is_raw_original
    data["raw_sidecar_path"] = work.raw_sidecar_path
    data["raw_processed_variant_id"] = work.raw_processed_variant_id

    return data
